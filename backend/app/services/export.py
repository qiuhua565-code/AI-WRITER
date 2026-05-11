import io
import re
from html import unescape
from html.parser import HTMLParser

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.models.task import Task
from app.models.segment import Segment

CHAPTER_TITLES = {
    "chapter_1": "第一章",
    "chapter_2": "第二章",
    "chapter_3": "第三章",
    "chapter_4": "第四章",
    "chapter_5": "第五章",
    "chapter_6": "第六章",
    "chapter_7": "第七章",
    "epilogue":  "尾声",
}

CHAPTER_ORDER = list(CHAPTER_TITLES.keys())

_HTML_TAG_RE = re.compile(r"<[^>]+>", re.DOTALL)
_BR_TAGS = re.compile(r"(?i)<br\s*/?>")


def _strip_all_tags_to_text(s: str) -> str:
    t = (s or "").strip()
    if not t:
        return ""
    if "<" in t and ">" in t:
        t = unescape(_HTML_TAG_RE.sub("", t))
    return t.strip()


class _InlineHtmlRunCollector(HTMLParser):
    """收集一行（或一段）里的纯文本片段及 b/strong、i/em 叠加态，供写入 docx run。"""

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.runs: list[tuple[str, bool, bool]] = []
        self._bold_depth = 0
        self._italic_depth = 0
        self._buf: list[str] = []

    def _flush(self) -> None:
        if not self._buf:
            return
        text = "".join(self._buf)
        self._buf.clear()
        if not text:
            return
        self.runs.append((text, self._bold_depth > 0, self._italic_depth > 0))

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        _ = attrs
        self._flush()
        t = tag.lower()
        if t in ("b", "strong"):
            self._bold_depth += 1
        elif t in ("i", "em"):
            self._italic_depth += 1

    def handle_endtag(self, tag: str) -> None:
        self._flush()
        t = tag.lower()
        if t in ("b", "strong") and self._bold_depth > 0:
            self._bold_depth -= 1
        elif t in ("i", "em") and self._italic_depth > 0:
            self._italic_depth -= 1

    def handle_data(self, data: str) -> None:
        self._buf.append(data)

    def close(self) -> None:
        self._flush()
        super().close()


def _add_paragraph_from_line(doc: Document, raw: str) -> None:
    """写入正文一行：保留 strong/em，否则退化为纯文本。"""
    line = (raw or "").rstrip("\r\n")
    if not line.strip():
        doc.add_paragraph()
        return
    frag = _BR_TAGS.sub(" ", line)
    if "<" not in frag or ">" not in frag:
        doc.add_paragraph(frag.strip())
        return
    collector = _InlineHtmlRunCollector()
    try:
        collector.feed(frag)
        collector.close()
    except Exception:
        plain = _strip_all_tags_to_text(frag)
        if plain:
            doc.add_paragraph(plain)
        else:
            doc.add_paragraph()
        return
    if not collector.runs:
        doc.add_paragraph()
        return
    p = doc.add_paragraph()
    for text, bold, italic in collector.runs:
        if not text:
            continue
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic


def build_docx(task: Task, segments: list[Segment]) -> bytes:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)

    # Title
    heading = doc.add_heading(f"回顾：{task.title}", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Declaration
    decl = doc.add_paragraph("本文根据真实社会事件改编，人物均已化名处理，如有雷同纯属巧合。")
    decl.runs[0].italic = True
    decl.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_paragraph()

    # Sort segments by their canonical order
    seg_map = {s.segment_type: s for s in segments}
    ordered = [seg_map[t] for t in CHAPTER_ORDER if t in seg_map and seg_map[t].content]

    for seg in ordered:
        title = CHAPTER_TITLES.get(seg.segment_type, seg.title or "")

        # Chapter divider + title
        divider = doc.add_paragraph(f"{'━' * 10}  {title}  {'━' * 10}")
        divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        # Body — 保留段落换行；行内 <strong>/<em> 等映射到 Word 加粗/斜体
        for line in seg.content.splitlines():
            _add_paragraph_from_line(doc, line)

        doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
